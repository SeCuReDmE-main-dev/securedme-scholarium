from __future__ import annotations

import re
import struct
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
GUIDE_DIR = ROOT / "docs" / "teach" / "guides"
OUTPUT_DIR = ROOT / "docs" / "teach" / "deliverables" / "guides"

PAPER = "F7F8FC"
INK = "10172F"
MUTED = "63708A"
BLUE = "2157EE"
CYAN = "23B8FF"
LINE = "DFE4EF"

GUIDES = {
    "STUDENT_GUIDE.md": {
        "output": "Scholarium_Teach_Student_Guide.docx",
        "role": "STUDENT GUIDE",
        "diagram": ROOT / "docs" / "teach" / "diagrams" / "02-data-lifecycle" / "scholarium-teach-data-lifecycle.png",
        "diagram_alt": "Scholarium Teach data lifecycle from purpose and consent to learner-controlled evidence and aggregate projections.",
        "diagram_caption": "System map: learner control remains present from observation through correction, sharing, and deletion.",
    },
    "TEACHER_GUIDE.md": {
        "output": "Scholarium_Teach_Teacher_Guide.docx",
        "role": "TEACHER GUIDE",
        "diagram": ROOT / "docs" / "teach" / "diagrams" / "04-assistant-authority" / "assistant-cooperation-and-authority-boundaries.png",
        "diagram_alt": "Assistant authority map showing consent-bound projections from learner assistant to teacher and aggregate-only administration.",
        "diagram_caption": "System map: assistants cooperate through bounded projections; the private learner graph is never exported.",
    },
    "ADMINISTRATOR_GUIDE.md": {
        "output": "Scholarium_Teach_Administrator_Guide.docx",
        "role": "ADMINISTRATOR GUIDE",
        "diagram": ROOT / "docs" / "teach" / "diagrams" / "05-security-boundary" / "gate5-and-transitional-security-boundary.png",
        "diagram_alt": "Gate5 and transitional security boundary from authenticated browser access to private workers, signed receipts, and technical observability.",
        "diagram_caption": "System map: Gate5 fails closed before execution and Bouncy Castle signs content-free terminal receipt digests.",
    },
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=130, start=170, bottom=130, end=170) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document, role: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"SCHOLARIUM TEACH  /  {role}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), CYAN)
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("PUBLIC PRE-ALPHA  |  SYNTHETIC DATA ONLY  |  ")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(fp)


def add_inline_markup(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string("6F42FF")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_status_callout(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "EAF9FF")
    set_cell_margins(cell)
    for index, line in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_inline_markup(p, line)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_title_block(doc: Document, title: str, role: str, metadata: list[str]) -> None:
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(10)
    run = eyebrow.add_run("SECUREDME EDUCATION  /  PUBLIC PRE-ALPHA")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(CYAN)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    title_run = title_p.add_run(title)
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor.from_string(INK)

    role_p = doc.add_paragraph()
    role_p.paragraph_format.space_after = Pt(18)
    role_run = role_p.add_run(role)
    role_run.bold = True
    role_run.font.name = "Calibri"
    role_run.font.size = Pt(10)
    role_run.font.color.rgb = RGBColor.from_string(BLUE)
    add_status_callout(doc, metadata)


def parse_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    paragraph_buffer: list[str] = []

    def flush() -> None:
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline_markup(p, " ".join(paragraph_buffer))
            paragraph_buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if stripped.startswith("# ") or stripped.startswith("**Edition:") or stripped.startswith("**Data boundary:") or stripped.startswith("**Operational boundary:") or stripped.startswith("**Governance boundary:"):
            continue
        if stripped.startswith("## "):
            flush()
            doc.add_heading(stripped[3:], level=1)
            continue
        if stripped.startswith("### "):
            flush()
            doc.add_heading(stripped[4:], level=2)
            continue
        if stripped.startswith("- "):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, stripped[2:])
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush()
            p = doc.add_paragraph(style="List Number")
            add_inline_markup(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue
        paragraph_buffer.append(stripped.rstrip("  "))
    flush()


def add_picture_with_alt(doc: Document, path: Path, alt_text: str, caption: str) -> None:
    heading = doc.add_heading("System Map", level=1)
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True
    with path.open("rb") as stream:
        signature = stream.read(24)
    if signature[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError(f"Expected PNG diagram: {path}")
    pixel_width, pixel_height = struct.unpack(">II", signature[16:24])
    width_inches = min(5.7, 7.4 * pixel_width / pixel_height)
    height_inches = width_inches * pixel_height / pixel_width
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches), height=Inches(height_inches))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", "Scholarium Teach system map")
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(6)
    for cr in cp.runs:
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor.from_string(MUTED)


def build_guide(source_name: str, config: dict[str, object]) -> Path:
    source_path = GUIDE_DIR / source_name
    text = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].removeprefix("# ").strip()
    metadata = [
        re.sub(r"\s{2,}", " ", line).strip()
        for line in lines[1:5]
        if line.strip().startswith("**")
    ]

    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, str(config["role"]))
    add_title_block(doc, title, str(config["role"]), metadata)
    parse_markdown(doc, text)
    add_picture_with_alt(
        doc,
        Path(config["diagram"]),
        str(config["diagram_alt"]),
        str(config["diagram_caption"]),
    )
    doc.core_properties.title = title
    doc.core_properties.subject = "Scholarium Teach role guide"
    doc.core_properties.author = "SecuredMe Education"
    doc.core_properties.keywords = "Scholarium Teach, education, pre-alpha, synthetic data"
    doc.core_properties.comments = "Generated from repository-controlled Markdown sources."

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / str(config["output"])
    doc.save(output)
    return output


def main() -> int:
    outputs = [build_guide(source, config) for source, config in GUIDES.items()]
    for output in outputs:
        print(output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
