from __future__ import annotations

import csv
import json
import re
import subprocess
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
TEACH = ROOT / "docs" / "teach"
OUTPUT = TEACH / "deliverables" / "recovery"
ASSETS = OUTPUT / "assets"

INK = "172036"
BLUE = "1E5AA8"
TEAL = "008B87"
CORAL = "DB5A42"
GREEN = "2F7D4A"
MUTED = "5D687A"
LINE = "D9E1EA"
PALE_BLUE = "EAF2FB"
PALE_TEAL = "E6F5F3"
PALE_CORAL = "FCEDE9"
PALE_GREEN = "EAF5EE"

GATES = {
    "G1": (1, 12), "G2": (13, 24), "G3": (25, 36), "G4": (37, 48),
    "G5": (49, 60), "G6": (61, 72), "G7": (73, 84), "G8": (85, 96),
    "G9": (97, 108), "G10": (109, 120), "G11": (121, 132),
    "G12": (133, 144), "G13": (145, 156), "G14": (157, 163),
}


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def read_actions() -> list[dict[str, str]]:
    with (TEACH / "ACTION_MATRIX.csv").open(encoding="utf-8-sig", newline="") as stream:
        rows = list(csv.DictReader(stream))
    if len(rows) != 163:
        raise ValueError(f"Expected 163 actions, received {len(rows)}")
    ids = [int(row["action_id"]) for row in rows]
    if ids != list(range(1, 164)):
        raise ValueError("Action identifiers must cover 001-163 exactly once")
    return rows


def parse_idea_registry() -> list[dict[str, str]]:
    lines = (TEACH / "IDEA_REGISTRY.md").read_text(encoding="utf-8").splitlines()
    rows = []
    for line in lines:
        if not line.startswith("|") or "---" in line or line.startswith("| Idea"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if len(cells) == 5:
            rows.append(dict(zip(("idea", "source", "classification", "decision", "destination"), cells)))
    return rows


def gate_for(action_id: int) -> str:
    return next(gate for gate, (start, end) in GATES.items() if start <= action_id <= end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in (
        ("Heading 1", 18, BLUE, 18, 8),
        ("Heading 2", 13.5, TEAL, 14, 6),
        ("Heading 3", 11.5, BLUE, 10, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run("SCHOLARIUM TEACH  /  RECOVERY MASTER")
        run.bold = True
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(BLUE)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = footer.add_run("PUBLIC PRE-ALPHA  |  SYNTHETIC DATA ONLY  |  ")
        fr.font.name = "Aptos"
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = RGBColor.from_string(MUTED)
        add_page_field(footer)


def add_title(doc: Document, state: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("SECUREDME EDUCATION  /  SYSTEM RECOVERY")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CORAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Scholarium + Teach")
    r.font.name = "Georgia"
    r.font.size = Pt(31)
    r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Dossier maître de récupération")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    add_callout(doc, [
        f"État vérifié: {state['programCompletion']['completed']}/163 actions terminées ({state['programCompletion']['percent']:.2f} %).",
        "Deux blocages externes restent visibles: validation Loi 25 et authentification interactive du tunnel VS Code.",
        "Ce document relie l'idée initiale, les décisions, l'implantation, les preuves et la reprise opérationnelle.",
    ], PALE_BLUE)
    p = doc.add_paragraph("Version du 16 juillet 2026  |  Source de vérité: docs/teach/EXECUTION_STATE.json")
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_callout(doc: Document, lines: list[str], fill: str = PALE_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    for index, line in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(9.5)
        if index == 0:
            run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def make_charts(actions: list[dict[str, str]]) -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    font_root = Path("C:/Windows/Fonts")
    regular = ImageFont.truetype(str(font_root / "segoeui.ttf"), 30)
    small = ImageFont.truetype(str(font_root / "segoeui.ttf"), 23)
    bold = ImageFont.truetype(str(font_root / "segoeuib.ttf"), 34)
    white = "#FFFFFF"
    status_counts = Counter(row["status"] for row in actions)
    status_path = ASSETS / "action-status.png"
    image = Image.new("RGB", (1500, 390), white)
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Programme de 163 actions: état vérifié", fill=f"#{INK}", font=bold)
    left = 60
    top = 145
    available = 1380
    colors = {"completed": "#2F7D4A", "blocked": "#DB5A42", "planned": "#1E5AA8", "in_progress": "#008B87"}
    labels = {"completed": "Terminées", "blocked": "Bloquées", "planned": "Planifiées", "in_progress": "En cours"}
    legend_x = 60
    for status in ("completed", "blocked", "planned", "in_progress"):
        count = status_counts.get(status, 0)
        if count:
            width = round(available * count / 163)
            draw.rounded_rectangle((left, top, left + width, top + 90), radius=10, fill=colors[status])
            if width > 60:
                text = str(count)
                box = draw.textbbox((0, 0), text, font=bold)
                draw.text((left + (width - (box[2] - box[0])) / 2, top + 22), text, fill=white, font=bold)
            left += width
            draw.rectangle((legend_x, 285, legend_x + 28, 313), fill=colors[status])
            draw.text((legend_x + 40, 280), f"{labels[status]}: {count}", fill=f"#{MUTED}", font=small)
            legend_x += 300
    image.save(status_path)

    gate_path = ASSETS / "gate-completion.png"
    gate_values = []
    for gate, (start, end) in GATES.items():
        subset = [row for row in actions if start <= int(row["action_id"]) <= end]
        gate_values.append(100 * sum(row["status"] == "completed" for row in subset) / len(subset))
    image = Image.new("RGB", (1500, 780), white)
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Complétion par porte d'exécution", fill=f"#{INK}", font=bold)
    baseline = 650
    chart_top = 130
    bar_width = 65
    gap = 31
    for index, (gate, value) in enumerate(zip(GATES, gate_values)):
        x = 70 + index * (bar_width + gap)
        height = round((baseline - chart_top) * value / 100)
        color = f"#{TEAL}" if value == 100 else f"#{CORAL}"
        draw.rounded_rectangle((x, baseline - height, x + bar_width, baseline), radius=7, fill=color)
        label = f"{value:.0f}%"
        box = draw.textbbox((0, 0), label, font=small)
        draw.text((x + (bar_width - (box[2] - box[0])) / 2, baseline - height - 34), label, fill=f"#{INK}", font=small)
        gate_box = draw.textbbox((0, 0), gate, font=small)
        draw.text((x + (bar_width - (gate_box[2] - gate_box[0])) / 2, baseline + 16), gate, fill=f"#{MUTED}", font=small)
    draw.line((60, baseline, 1430, baseline), fill=f"#{LINE}", width=2)
    image.save(gate_path)

    timeline_path = ASSETS / "project-timeline.png"
    events = [
        "Cours\nd'espagnol", "Mémoire et\nmaîtrise", "Accessibilité\nedge-case-first", "Assistants et\nforces",
        "Social et\nportfolio", "Média et\nplugins", "Gate5 et\npreuves", "Stitch, QA et\ndéploiement",
    ]
    image = Image.new("RGB", (1500, 540), white)
    draw = ImageDraw.Draw(image)
    draw.text((60, 30), "Évolution conservée: du besoin réel à la suite éducative", fill=f"#{INK}", font=bold)
    y = 280
    draw.line((95, y, 1405, y), fill="#AAB6C4", width=6)
    palette = (CORAL, BLUE, TEAL, GREEN, CORAL, BLUE, TEAL, GREEN)
    for index, (label, color) in enumerate(zip(events, palette)):
        x = 110 + index * 182
        draw.ellipse((x - 22, y - 22, x + 22, y + 22), fill=f"#{color}", outline=white, width=4)
        lines = label.split("\n")
        label_y = 145 if index % 2 == 0 else 340
        for line_index, line in enumerate(lines):
            box = draw.textbbox((0, 0), line, font=small)
            draw.text((x - (box[2] - box[0]) / 2, label_y + line_index * 30), line, fill=f"#{MUTED}", font=small)
    image.save(timeline_path)
    return {"status": status_path, "gates": gate_path, "timeline": timeline_path}


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    output = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    output.extend("| " + " | ".join(value.replace("|", "\\|").replace("\n", " ") for value in row) + " |" for row in rows)
    return "\n".join(output)


def build_markdown(actions, ideas, state, runtime, deliverables) -> str:
    counts = Counter(row["status"] for row in actions)
    gate_rows = []
    for gate, (start, end) in GATES.items():
        subset = [row for row in actions if start <= int(row["action_id"]) <= end]
        gate_rows.append([gate, f"{start:03d}-{end:03d}", str(sum(row["status"] == "completed" for row in subset)), str(sum(row["status"] == "blocked" for row in subset)), str(len(subset))])
    action_rows = [[row["action_id"], gate_for(int(row["action_id"])), row["action"], row["skills"], row["plugins"], row["status"], row["evidence"]] for row in actions]
    idea_rows = [[item["idea"], item["source"], item["classification"], item["decision"], item["destination"]] for item in ideas]
    production_rows = [[item["url"], str(item["status"]), str(item["bytes"]), item["contentCheck"]] for item in runtime["production"]]
    return f"""# Scholarium + Teach - Dossier maître de récupération

Date: 2026-07-16  
Statut: public pre-alpha; données synthétiques seulement  
Source de vérité: `docs/teach/EXECUTION_STATE.json`

## Résumé exécutif

- **{counts['completed']}/163 actions sont terminées ({state['programCompletion']['percent']:.2f} %).**
- **Deux actions demeurent bloquées par des conditions externes:** 048 pour la validation Loi 25 et 157 pour l'authentification interactive du tunnel VS Code.
- **L'application est active:** landing, `/app`, `/teach`, santé API et leçon espagnole ont répondu HTTP 200.
- **La PR #2 est fusionnée:** {runtime['pullRequest']['url']}.
- **Le corpus G5 est fermé:** 125 cartes, 125 classifications Synthia, 125 paquets OpenIE et trois magasins de graphe durables.

Ce dossier permet de récupérer le projet sans relire la conversation. Il sépare l'idée initiale, les décisions, ce qui est réellement implanté, les preuves et ce qui demeure externe.

![État des actions](assets/action-status.png)

## 1. Point de départ: une vraie séance d'espagnol

Le projet est parti d'un besoin concret: apprendre à tenir une conversation quotidienne simple en espagnol au moyen de séances réelles d'environ une heure, pratiques, orales, visuelles, progressives et fondées sur la maîtrise.

Quatre interactions ont servi de premier laboratoire:

1. `Hola, ¿cómo estás?` -> `Muy bien, ¿y tú?`
2. `¿Cómo te llamas?` -> `Me llamo Jean-Sébastien.`
3. `¿De dónde eres?` -> `Soy de Montreal.`
4. `¿Cuántos años tienes?` -> `Tengo treinta y ocho años.`

Les erreurs observées ont établi les premiers contrats: une réponse doit correspondre à la question; une lecture ou un « OK » n'est pas une preuve de maîtrise; la phrase peut être segmentée puis reconstruite; les rappels doivent être immédiats, différés et contextuels; une séance doit reprendre exactement après interruption.

![Évolution du projet](assets/project-timeline.png)

## 2. Idée centrale et invariants

**Un élève n'est pas résumé par ses notes.** Une note est un signal parmi d'autres. Scholarium Teach relie performance, processus d'apprentissage, talents, passions, projets, sport, musique, créativité, persévérance et contribution sociale.

L'exemple fondateur reste: `D en math -> réussite au soccer -> stratégie spatiale`. L'assistant peut proposer un pont entre lecture du jeu, angles, trajectoires, anticipation et géométrie. L'élève conserve le droit d'accepter, reformuler, contester, expirer ou supprimer cette interprétation.

Invariants conservés:

- empowerment, dignité et pluralité des formes d'intelligence;
- intégrité et respect plutôt qu'une conception restriction-first;
- expression honnête des difficultés sans positivité forcée;
- provenance, contradiction, incertitude et correction visibles;
- contrôle utilisateur et séparation des autorités;
- accessibilité edge-case-first;
- aucune surveillance cachée, écoute passive, décision disciplinaire autonome ou diagnostic clinique.

## 3. Évolution du système

### Mémoire pédagogique et maîtrise

Cours, modules, leçons, notions, questions, réponses, indices, niveaux d'aide, rappels et checkpoints forment une boucle durable. La maîtrise exige une réponse non assistée, liée à la bonne question, rappelée après délai et transférée dans un nouveau contexte.

### AlgoQuest et assistants

AlgoQuest sert de porte d'entrée éducative et d'outbox durable. L'assistant étudiant conserve le graphe privé. L'enseignant reçoit une projection pédagogique bornée; l'administration reçoit uniquement des agrégats respectant le seuil de cohorte.

### Accessibilité edge-case-first

Sept profils combinables ont été implantés: sourd/signé disponible, non verbal, Autism Calm, Tourette Safe, ADHD Sprint, Dyslexia Reading et Dyspraxia Motor. Les parcours essentiels fonctionnent sans voix, souris ou animation obligatoire.

### Social, portfolio et analytique

Le système comprend capsules de croissance, fils de projets, cercles, reconnaissances et récapitulatifs. Les statistiques sont utiles à l'élève et pratiques pour l'enseignant sans produire un score unique de valeur personnelle.

### Média et plugins éducatifs

Les médias sont déclenchés explicitement depuis une leçon, un fil, un projet ou une réussite. Les contrats plafonnent à trois vidéos quotidiennes de cinq minutes et cinq balados de trente minutes, avec confirmation distincte avant publication.

Six plugins éducatifs stateless possèdent des starters complets `AGENTS.md`, `SOUL.md` et `USER.md`. Aucun plugin ne possède de `MEMORY.md`; la mémoire est centralisée par le gateway.

### Synthia, MemoryLake et HippoRAG

Synthia préserve `I -> I_system^S -> H_lex -> G_lex -> I_lexicon`. MemoryLake fournit `index_records`; HippoRAG demeure limité à `retrieve_dpr`; Codex ou Gemini produit le langage utilisateur. Synthia trace les termes, sources, corrections et incertitudes sans devenir autorité pédagogique, scientifique, juridique ou taxonomique.

### FNP-QNN, FfeD, Gate5 et géométrie

Pluginpack a passé 337 tests, 41/41 contrôles d'intégrité et 12/12 contrôles doctor. FNP-QNN exécute cinq plugins sous allowlist et poids versionnés. Gate5 contrôle capacité, consentement, expiration, replay et preuve. La géométrie quasicristalline fournit adressage, projection, replay et diff; elle ne remplace pas la cryptographie et n'entre jamais dans KDF, AAD, coffre ou matériel de clé.

## 4. Ce qui a réellement été construit

- boucle espagnole complète avec checkpoint et scénario d'environ une heure;
- sept profils d'accessibilité, clavier, ARIA et contrôles de mouvement/son/densité;
- assistants, objectifs hebdomadaires, interventions silencieuses et miroir des forces;
- portfolio, projets, cercles, reconnaissances, récapitulatifs et tableaux analytiques;
- contrats QuaNTecH-ViD, quotas et manifestes signés contre altération et replay;
- six prefabs éducatifs avec starters Codex/Gemini et gateway mémoire central;
- corpus de 125 sources, Synthia, OpenIE et graphes approuvé/préparation/quarantaine;
- Gate5, FNP-QNN, pluginpack, Bouncy Castle transitoire et adressage structurel;
- identité Stitch, navigation Teach, Sources et Administration;
- guides élève/enseignant/administrateur, deux présentations et dossier corporatif;
- tests, captures, migrations, OpenAPI, sauvegarde/restauration, Datadog et déploiement Cloudflare.

## 5. État quantitatif des 163 actions

{markdown_table(['Porte', 'Actions', 'Terminées', 'Bloquées', 'Total'], gate_rows)}

![Complétion par porte](assets/gate-completion.png)

La frontière d'exécution reste `156`. La dette historique 049-057 a été fermée sans rembobiner les groupes déjà validés.

## 6. État opérationnel vérifié

{markdown_table(['Surface', 'HTTP', 'Octets', 'Contrôle'], production_rows)}

- Branche de travail lors de la capture: `{runtime['branch']}`.
- Commit de référence avant le dossier: `{runtime['headBeforeRecoveryCommit']}`.
- PR #2: `{runtime['pullRequest']['state']}`, fusionnée le `{runtime['pullRequest']['mergedAt']}`.
- Check Datadog de la PR: `{runtime['pullRequest']['datadogCheck']}`.
- Le fichier utilisateur `assets/desing/ChatGPT Installer (1).exe` est exclu de la livraison Git.

## 7. Reprise exacte

Fichiers faisant autorité:

1. `docs/teach/EXECUTION_STATE.json` - pointeur de reprise.
2. `docs/teach/ACTION_MATRIX.csv` - statut et preuve de chaque action.
3. `docs/teach/EVIDENCE_REGISTER.md` - index des preuves.
4. `docs/teach/IDEA_REGISTRY.md` - idées conservées, reportées ou rejetées.
5. `docs/teach/corpus/classification-run.json` - reçu G5.

Procédure après reboot:

1. Lire `activeAction` dans `EXECUTION_STATE.json`; ne jamais choisir la plus ancienne ligne incomplète de la matrice.
2. Vérifier `git status --short` et préserver tous les fichiers utilisateur non suivis.
3. Lancer les tests ciblés de l'action active.
4. Consigner une preuve directe avant de modifier un statut.
5. Ne jamais faire reculer `executionFrontier` pour régler une dette historique.

Blocages restant à fermer:

- **048:** obtenir une revue Loi 25 qualifiée avant toute donnée réelle d'élève; conserver la porte France/UE séparée.
- **157:** terminer l'authentification interactive du tunnel VS Code; le développement local reste fonctionnel.

## 8. Corpus G5 et limites explicites

- 125 cartes uniques et 125 classifications Synthia réelles.
- 66 URL accessibles, 54 restreintes, 2 erreurs HTTP et 3 vérifications réseau non résolues.
- 125 paquets OpenIE; aucune contradiction sémantique inventée.
- `approved=0`, `preparation=125`, `quarantine=0`.
- Les graphes sont opérationnels; leur partition ne constitue jamais une certification.

## 9. Index des livrables

""" + "\n".join(f"- `{path}`" for path in deliverables) + f"""

## Annexe A - Registre complet des idées

{markdown_table(['Idée', 'Source', 'Classe', 'Décision', 'Destination'], idea_rows)}

## Annexe B - Registre complet des 163 actions

{markdown_table(['ID', 'Porte', 'Action', 'Skills', 'Plugins', 'Statut', 'Preuve'], action_rows)}

## Annexe C - Carte de provenance

- Idée originale: `docs/IDEA_CAPTURE_2026-07-15_scholarium_spanish_course.md`.
- Rapport émergent: `docs/idea-captures/2026-07-15_scholarium-teach-spanish-emergent-design.md`.
- Brief opérationnel: `docs/idea-captures/2026-07-15_scholarium-teach_operational-brief.md`.
- Architecture: `docs/teach/ARCHITECTURE.md`.
- Plan accepté: `docs/teach/PLAN_163_ACTIONS.md`.
- Preuves: `docs/teach/evidence/` et `docs/teach/EVIDENCE_REGISTER.md`.
- État en ligne: `docs/teach/evidence/recovery-master-runtime.json`.

## Annexe D - Glossaire

- **Scholarium:** plateforme éducative, sociale et de portfolio.
- **Teach:** domaine adaptatif de l'application Scholarium existante.
- **AlgoQuest:** porte d'entrée et échange éducatif durable.
- **Synthia:** gardienne de lexiques et de traçabilité sous autorité humaine.
- **MemoryLake:** stockage central approuvé par contrat `index_records`.
- **HippoRAG:** récupération seulement par `retrieve_dpr`.
- **Gate5:** courtier de capacités, preuves, consentements, replay et audit.
- **FNP-QNN:** moteur expérimental de signaux plithogéniques derrière Gate5.
- **QuaNTecH-ViD:** worker média privé, explicitement déclenché et plafonné.
- **V.O.T Guardian/Tenebris:** enveloppe éphémère étudiée; aucune écoute passive ni analyse biométrique automatique.

## Conclusion

Le projet n'a pas disparu dans la journée de travail. Il est passé d'une séance d'espagnol à une suite éducative pre-alpha active, documentée et testée. La fermeture honnête est `161/163`: tout le travail interne prévu est fermé; les deux dernières actions dépendent d'une validation juridique externe et d'une authentification interactive.
"""


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    for run in caption_p.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_status_table(doc, actions) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("Porte", "Actions", "Terminées", "Bloquées", "Total")
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for index, (gate, (start, end)) in enumerate(GATES.items()):
        subset = [row for row in actions if start <= int(row["action_id"]) <= end]
        values = (gate, f"{start:03d}-{end:03d}", str(sum(row["status"] == "completed" for row in subset)), str(sum(row["status"] == "blocked" for row in subset)), str(len(subset)))
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value
            if index % 2:
                set_cell_shading(cell, "F4F7FA")
    set_table_widths(table, [0.7, 1.2, 1.35, 1.2, 0.8])


def add_production_table(doc, runtime) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ("Surface", "HTTP", "Contrôle")):
        cell.text = header
        set_cell_shading(cell, TEAL)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for item in runtime["production"]:
        row = table.add_row()
        for cell, value in zip(row.cells, (item["url"], str(item["status"]), item["contentCheck"])):
            cell.text = value
    set_table_widths(table, [3.8, 0.7, 2.25])


def add_action_appendix(doc: Document, actions) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    doc.add_heading("Annexe B - Registre complet des 163 actions", level=1)
    current_gate = None
    for action in actions:
        gate = gate_for(int(action["action_id"]))
        if gate != current_gate:
            doc.add_heading(f"{gate} - Actions {GATES[gate][0]:03d} à {GATES[gate][1]:03d}", level=2)
            current_gate = gate
        header = doc.add_paragraph()
        header.paragraph_format.keep_with_next = True
        header.paragraph_format.space_before = Pt(6)
        header.paragraph_format.space_after = Pt(2)
        run = header.add_run(f"{action['action_id']}  |  {action['status'].upper()}  |  {action['skills']}  |  {action['plugins']}")
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(GREEN if action["status"] == "completed" else CORAL)
        body = doc.add_paragraph(action["action"])
        body.paragraph_format.keep_with_next = True
        body.paragraph_format.space_after = Pt(2)
        for run in body.runs:
            run.font.size = Pt(9)
        evidence = doc.add_paragraph(f"Preuve: {action['evidence']}")
        evidence.paragraph_format.space_after = Pt(5)
        border = evidence._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), LINE)
        p_bdr.append(bottom)
        border.append(p_bdr)
        for run in evidence.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def build_docx(actions, ideas, state, runtime, deliverables, charts, path: Path) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, state)
    doc.add_heading("Résumé exécutif", level=1)
    add_bullets(doc, [
        f"{state['programCompletion']['completed']}/163 actions terminées ({state['programCompletion']['percent']:.2f} %).",
        "Deux actions bloquées par des conditions externes: 048 (Loi 25) et 157 (tunnel VS Code).",
        "Cinq surfaces publiques ont répondu HTTP 200 pendant la capture.",
        "La PR #2 est fusionnée et son check Datadog a réussi.",
        "G5 contient 125 cartes, 125 reçus Synthia, 125 paquets OpenIE et trois magasins durables.",
    ])
    add_picture(doc, charts["status"], 6.5, "État recalculé directement depuis ACTION_MATRIX.csv.")

    doc.add_heading("1. Point de départ: une vraie séance d'espagnol", level=1)
    doc.add_paragraph("Le projet est parti d'un besoin concret: apprendre à tenir une conversation quotidienne simple en espagnol au moyen de séances réelles d'environ une heure, pratiques, orales, visuelles, progressives et fondées sur la maîtrise.")
    interactions = [
        "Hola, ¿cómo estás? -> Muy bien, ¿y tú?",
        "¿Cómo te llamas? -> Me llamo Jean-Sébastien.",
        "¿De dónde eres? -> Soy de Montreal.",
        "¿Cuántos años tienes? -> Tengo treinta y ocho años.",
    ]
    for item in interactions:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("Les erreurs observées ont établi les premiers contrats: la réponse doit correspondre à la question; lire ou dire « OK » ne prouve pas la maîtrise; la phrase peut être segmentée puis reconstruite; les rappels doivent être différés et contextuels; la séance doit reprendre exactement après interruption.")
    add_picture(doc, charts["timeline"], 6.55, "Du laboratoire espagnol à la suite éducative complète.")

    doc.add_heading("2. Idée centrale et invariants", level=1)
    add_callout(doc, ["Un élève n'est pas résumé par ses notes.", "Une note est un signal. Les talents, passions, projets, efforts et contributions sont aussi des signaux."], PALE_CORAL)
    doc.add_paragraph("L'exemple fondateur reste: D en math -> réussite au soccer -> stratégie spatiale. L'assistant propose un pont entre lecture du jeu, angles, trajectoires, anticipation et géométrie. L'élève peut accepter, reformuler, contester, expirer ou supprimer l'interprétation.")
    add_bullets(doc, [
        "Empowerment, dignité et pluralité des formes d'intelligence.",
        "Contrats d'intégrité et de respect plutôt qu'une conception restriction-first.",
        "Expression honnête des difficultés sans positivité forcée.",
        "Provenance, contradiction, incertitude et correction visibles.",
        "Accessibilité edge-case-first et contrôle utilisateur.",
        "Aucune surveillance cachée, écoute passive, discipline autonome ou prétention clinique.",
    ])

    doc.add_heading("3. Évolution du système", level=1)
    sections = [
        ("Mémoire pédagogique et maîtrise", "Cours, modules, leçons, notions, questions, indices, niveaux d'aide, rappels et checkpoints forment une boucle durable. La maîtrise exige une réponse non assistée, liée à la bonne question, rappelée après délai et transférée."),
        ("AlgoQuest et assistants", "AlgoQuest est une outbox éducative durable. Le graphe étudiant reste privé; l'enseignant reçoit une projection bornée et l'administration uniquement des agrégats respectant le seuil de cohorte."),
        ("Accessibilité edge-case-first", "Sept profils combinables couvrent sourd/signé disponible, non verbal, Autism Calm, Tourette Safe, ADHD Sprint, Dyslexia Reading et Dyspraxia Motor. Les parcours essentiels fonctionnent sans voix, souris ou animation obligatoire."),
        ("Social, portfolio et analytique", "Capsules de croissance, projets, cercles, reconnaissances et récapitulatifs amplifient la progression sans produire un score unique de valeur personnelle."),
        ("Média et plugins éducatifs", "Les médias sont déclenchés explicitement depuis une source. Trois vidéos de cinq minutes et cinq balados de trente minutes constituent les plafonds initiaux. Six plugins stateless disposent de AGENTS.md, SOUL.md et USER.md; aucun n'utilise MEMORY.md."),
        ("Synthia, MemoryLake et HippoRAG", "Synthia préserve I -> I_system^S -> H_lex -> G_lex -> I_lexicon. MemoryLake indexe, HippoRAG récupère seulement, Codex ou Gemini génère. La décision reste humaine."),
        ("FNP-QNN, FfeD, Gate5 et géométrie", "Gate5 contrôle capacité, consentement, expiration, replay et preuve. FNP-QNN reste derrière le gateway. La géométrie sert à adresser, projeter, rejouer et comparer; elle ne remplace pas la cryptographie."),
    ]
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    doc.add_heading("4. Ce qui a réellement été construit", level=1)
    add_bullets(doc, [
        "Boucle espagnole complète avec checkpoint et scénario d'environ une heure.",
        "Sept profils d'accessibilité, clavier, ARIA et contrôles sensoriels.",
        "Assistants, objectifs hebdomadaires, intervention silencieuse et miroir des forces.",
        "Portfolio, projets, cercles, reconnaissances, récapitulatifs et tableaux analytiques.",
        "QuaNTecH-ViD, quotas, manifestes signés et protection contre altération/replay.",
        "Six prefabs éducatifs et gateway mémoire central sans mémoire intra-plugin.",
        "125 sources, Synthia, OpenIE et graphes approuvé/préparation/quarantaine.",
        "Pluginpack, FNP-QNN, Gate5, contrôle Bouncy Castle transitoire et adressage structurel.",
        "Identité Stitch, Teach, Sources, Administration, guides, présentations et dossier corporatif.",
        "Tests, captures, migrations, OpenAPI, sauvegarde/restauration, Datadog et Cloudflare.",
    ])

    doc.add_heading("5. État quantitatif des 163 actions", level=1)
    add_status_table(doc, actions)
    add_picture(doc, charts["gates"], 6.55, "Toutes les portes sont fermées sauf les conditions externes de G4 et G14.")
    add_callout(doc, ["Frontière d'exécution: 156.", "La fermeture historique 049-057 ne déplace pas la frontière et ne redémarre aucun groupe validé."], PALE_BLUE)

    doc.add_heading("6. État opérationnel vérifié", level=1)
    add_production_table(doc, runtime)
    add_bullets(doc, [
        f"Branche lors de la capture: {runtime['branch']}.",
        f"Commit de référence avant dossier: {runtime['headBeforeRecoveryCommit']}.",
        f"PR #2: {runtime['pullRequest']['state']}, fusionnée le {runtime['pullRequest']['mergedAt']}.",
        f"Check Datadog: {runtime['pullRequest']['datadogCheck']}.",
        "Le fichier utilisateur assets/desing/ChatGPT Installer (1).exe est exclu de Git.",
    ])

    doc.add_heading("7. Reprise exacte", level=1)
    doc.add_paragraph("Ordre des sources de vérité:")
    for item in [
        "docs/teach/EXECUTION_STATE.json - pointeur de reprise.",
        "docs/teach/ACTION_MATRIX.csv - statut de chaque action.",
        "docs/teach/EVIDENCE_REGISTER.md - index des preuves.",
        "docs/teach/IDEA_REGISTRY.md - conservation des idées.",
        "docs/teach/corpus/classification-run.json - reçu G5.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("Après reboot:")
    add_bullets(doc, [
        "Lire activeAction dans EXECUTION_STATE.json; ne jamais reprendre à la plus ancienne ligne incomplète.",
        "Vérifier git status --short et préserver tous les fichiers utilisateur.",
        "Exécuter les tests ciblés avant de modifier un statut.",
        "Consigner une preuve directe et ne jamais faire reculer executionFrontier.",
    ])
    add_callout(doc, [
        "Action 048: revue Loi 25 qualifiée requise avant toute donnée réelle d'élève.",
        "Action 157: authentification interactive du tunnel VS Code; le développement local fonctionne.",
    ], PALE_CORAL)

    doc.add_heading("8. Corpus G5 et limites explicites", level=1)
    add_bullets(doc, [
        "125 cartes uniques, 125 classifications Synthia et 125 paquets OpenIE.",
        "66 URL accessibles, 54 restreintes, 2 erreurs HTTP et 3 non résolues.",
        "approved=0, preparation=125, quarantine=0.",
        "Aucune contradiction sémantique, causalité ou certification n'est inventée depuis un titre.",
        "MemoryLake.index_records stocke; HippoRAG.retrieve_dpr récupère; Codex ou Gemini génère.",
    ])

    doc.add_heading("9. Index des livrables", level=1)
    for item in deliverables:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Annexe A - Registre complet des idées", level=1)
    for index, idea in enumerate(ideas, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{index:02d}. {idea['idea']}")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
        doc.add_paragraph(f"Source: {idea['source']}  |  Classe: {idea['classification']}  |  Décision: {idea['decision']}")
        p = doc.add_paragraph(f"Destination: {idea['destination']}")
        p.paragraph_format.space_after = Pt(8)

    add_action_appendix(doc, actions)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    doc.add_heading("Annexe C - Provenance et glossaire", level=1)
    add_bullets(doc, [
        "Idée originale: docs/IDEA_CAPTURE_2026-07-15_scholarium_spanish_course.md.",
        "Rapport émergent: docs/idea-captures/2026-07-15_scholarium-teach-spanish-emergent-design.md.",
        "Brief opérationnel: docs/idea-captures/2026-07-15_scholarium-teach_operational-brief.md.",
        "Architecture: docs/teach/ARCHITECTURE.md.",
        "Plan accepté: docs/teach/PLAN_163_ACTIONS.md.",
        "État en ligne: docs/teach/evidence/recovery-master-runtime.json.",
    ])
    glossary = [
        ("Scholarium", "Plateforme éducative, sociale et de portfolio."),
        ("Teach", "Domaine adaptatif de l'application Scholarium existante."),
        ("AlgoQuest", "Porte d'entrée et échange éducatif durable."),
        ("Synthia", "Gardienne de lexiques et de traçabilité sous autorité humaine."),
        ("MemoryLake", "Stockage central par index_records."),
        ("HippoRAG", "Récupération seulement par retrieve_dpr."),
        ("Gate5", "Courtier de capacités, preuves, consentements, replay et audit."),
        ("FNP-QNN", "Moteur expérimental plithogénique derrière Gate5."),
        ("QuaNTecH-ViD", "Worker média privé, explicitement déclenché et plafonné."),
        ("V.O.T Guardian/Tenebris", "Enveloppe éphémère étudiée; aucune écoute passive ou route biométrique automatique."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        r = p.add_run(f"{term}: ")
        r.bold = True
        p.add_run(definition)
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Le projet est passé d'une séance d'espagnol à une suite éducative pre-alpha active, documentée et testée. La fermeture honnête est 161/163: tout le travail interne prévu est fermé; les deux dernières actions dépendent d'une validation juridique externe et d'une authentification interactive.")
    add_header_footer(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    actions = read_actions()
    ideas = parse_idea_registry()
    state = read_json(TEACH / "EXECUTION_STATE.json")
    runtime = read_json(TEACH / "evidence" / "recovery-master-runtime.json")
    counts = Counter(row["status"] for row in actions)
    expected = state["programCompletion"]
    if counts["completed"] != expected["completed"] or counts["blocked"] != expected["blocked"]:
        raise ValueError("Action matrix and execution state counts diverge")
    deliverables = sorted(str(path.relative_to(ROOT)).replace("\\", "/") for path in (TEACH / "deliverables").rglob("*") if path.is_file() and "recovery" not in path.parts)
    charts = make_charts(actions)
    markdown = build_markdown(actions, ideas, state, runtime, deliverables)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    md_path = OUTPUT / "Scholarium_Teach_Recovery_Master.md"
    docx_path = OUTPUT / "Scholarium_Teach_Recovery_Master.docx"
    md_path.write_text(markdown, encoding="utf-8", newline="\n")
    build_docx(actions, ideas, state, runtime, deliverables, charts, docx_path)
    print(json.dumps({"markdown": str(md_path), "docx": str(docx_path), "actions": len(actions), "ideas": len(ideas), "deliverables": len(deliverables)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
